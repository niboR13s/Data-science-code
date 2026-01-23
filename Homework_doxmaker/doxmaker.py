import os
import subprocess
import json
import datetime
from docx import Document
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH

# --- USER CONFIGURATION ---
STUDENT_INFO = {
    "name": "Robin Somers",
    "student_number": "1915369",
    "group": "IMAGINE",
    "course": "MNLE"
}

# Path to the specific Data Flavor project (relative or absolute path)
EXTERNAL_PROJECTS = [
    {
        "name": "Data Flavor Project/",
        "path": "Data_Project" 
    }
]
# --------------------------

def add_hyperlink(paragraph, url, text):
    """
    A helper function to add a clickable hyperlink to a paragraph.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    c = OxmlElement('w:color')
    c.set(qn('w:val'), "0000FF")
    rPr.append(c)

    u = OxmlElement('w:u')
    u.set(qn('w:val'), "single")
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink

def create_title_page(doc, title, subtitle="Homework Report"):
    """
    Creates a dedicated title page with student details.
    """
    # Title
    doc.add_heading(title, 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle
    p_sub = doc.add_paragraph(subtitle)
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.runs[0].italic = True
    
    # Spacing
    doc.add_paragraph("\n" * 5)
    
    # Student Info
    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    info_lines = [
        f"Student: {STUDENT_INFO['name']}",
        f"Student Number: {STUDENT_INFO['student_number']}",
        f"Group: {STUDENT_INFO['group']}",
        f"Course: {STUDENT_INFO['course']}",
        f"Date: {datetime.datetime.now().strftime('%Y-%m-%d')}"
    ]
    
    for line in info_lines:
        run = p_info.add_run(line + "\n")
        run.bold = True
        run.font.size = 160000 # Roughly 12.5pt (internal units)

    # Force page break after title page
    doc.add_page_break()

def get_git_remote_info():
    """
    Automatically detects the git remote URL and current branch.
    """
    try:
        remote_url = subprocess.check_output(["git", "config", "--get", "remote.origin.url"]).decode().strip()
        branch = subprocess.check_output(["git", "branch", "--show-current"]).decode().strip()

        if remote_url.startswith("git@"):
            remote_url = remote_url.replace(":", "/")
            remote_url = remote_url.replace("git@", "https://")
        
        if remote_url.endswith(".git"):
            remote_url = remote_url[:-4]

        return f"{remote_url}/blob/{branch}"

    except Exception as e:
        print(f"Warning: Could not auto-detect Git info ({e}). Using placeholder.")
        return "https://github.com/YOUR_USER/YOUR_REPO/blob/main"

def get_file_description(file_path):
    """
    Extracts a description based on file type.
    """
    _, ext = os.path.splitext(file_path)
    
    try:
        # HANDLING JUPYTER NOTEBOOKS
        if ext.lower() == '.ipynb':
            with open(file_path, "r", encoding="utf-8") as f:
                data = json.load(f)
                for cell in data.get('cells', []):
                    if cell.get('cell_type') == 'markdown':
                        source_content = cell.get('source', [])
                        if isinstance(source_content, list):
                            return "".join(source_content).strip()
                        return str(source_content).strip()
            return None

        # HANDLING REGULAR CODE FILES
        else:
            with open(file_path, "r", encoding="utf-8") as f:
                lines = f.readlines()
                comments = []
                for line in lines[:10]: 
                    stripped = line.strip()
                    if stripped.startswith(("#", "//", "/*")):
                        comments.append(stripped)
                
                if comments:
                    return " ".join(comments)

    except Exception as e:
        print(f"Error reading description for {file_path}: {e}")
    
    return None

def add_files_to_document(doc, file_list, github_base_url):
    """
    Helper to add a list of files to a document object.
    """
    for file_full_path in file_list:
        filename = os.path.basename(file_full_path)
        
        # 1. Heading (Filename)
        doc.add_heading(filename, level=2)

        # 2. Dynamic GitHub Link
        # Calculate relative path from the current working directory
        try:
            rel_path = os.path.relpath(file_full_path, start=os.getcwd())
            url_path = rel_path.replace("\\", "/")
            full_link = f"{github_base_url}/{url_path}"

            p_link = doc.add_paragraph()
            p_link.add_run("Link: ").bold = True
            add_hyperlink(p_link, full_link, "View on GitHub")
        except ValueError:
            # Fallback if file is on a different drive or path issue
            p_err = doc.add_paragraph("Could not generate relative link (External path).")

        # 3. Description
        desc_text = get_file_description(file_full_path)
        if desc_text:
            p_desc = doc.add_paragraph()
            p_desc.add_run("Description: ").bold = True
            if len(desc_text) > 500:
                desc_text = desc_text[:500] + "..."
            p_desc.add_run(desc_text).italic = True

def process_folder(folder_path, allowed_extensions):
    """
    Recursively collects all valid files in a folder and sorts them.
    Returns a sorted list of full file paths.
    """
    collected_files = []
    if not os.path.exists(folder_path):
        print(f"Warning: Path not found: {folder_path}")
        return collected_files

    for root, dirs, files in os.walk(folder_path):
        for file in files:
            _, ext = os.path.splitext(file)
            if ext.lower() in allowed_extensions:
                collected_files.append(os.path.join(root, file))
    
    # Sort files alphabetically
    collected_files.sort()
    return collected_files

def generate_weekly_reports():
    # --- CONFIGURATION ---
    base_folder = "Homework"
    output_folder = "Homework_doxmaker" 
    allowed_extensions = [".py", ".md", ".java", ".html", ".css", ".js", ".sql", ".ipynb"]
    
    github_base_url = get_git_remote_info()
    print(f"--- Detected Git URL: {github_base_url} ---")
    # ---------------------

    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    if not os.path.exists(base_folder):
        print(f"Error: Source folder '{base_folder}' not found.")
        return

    # Initialize Master Document
    master_doc = Document()
    create_title_page(master_doc, "Complete Homework Overview", subtitle="Master Portfolio")

    # Get list of weeks (folders) and sort them
    items = sorted(os.listdir(base_folder))

    for item in items:
        week_path = os.path.join(base_folder, item)

        if os.path.isdir(week_path):
            print(f"Processing folder: {item}...")
            
            # 1. Collect and Sort files for this week
            files_in_week = process_folder(week_path, allowed_extensions)
            
            if files_in_week:
                # --- GENERATE INDIVIDUAL WEEKLY REPORT ---
                weekly_doc = Document()
                create_title_page(weekly_doc, f"Homework: {item}")
                
                # Add Heading for the content
                weekly_doc.add_heading(f'Overview: {item}', level=1)
                
                # Add files
                add_files_to_document(weekly_doc, files_in_week, github_base_url)
                
                # Save Individual Report
                output_filename = f"Overview_{item}.docx"
                full_save_path = os.path.join(output_folder, output_filename)
                weekly_doc.save(full_save_path)
                print(f"  -> Saved individual report: {full_save_path}")

                # --- APPEND TO MASTER DOCUMENT ---
                # Add a Heading 1 for the week (This starts a new section)
                master_doc.add_heading(f"Week: {item}", level=1)
                
                # Add the same files to master doc
                add_files_to_document(master_doc, files_in_week, github_base_url)
                
                # Add a page break after each week to ensure separation
                master_doc.add_page_break()
                
            else:
                print(f"  -> No valid files found in {item}")

    # --- ADD EXTERNAL PROJECTS (DATA FLAVOR) ---
    print("Processing External Projects...")
    
    # Only add header if we have external projects
    if EXTERNAL_PROJECTS:
        master_doc.add_heading("External Projects", level=1)
        
        for project in EXTERNAL_PROJECTS:
            name = project['name']
            path = project['path']
            
            print(f"  -> Adding External Project: {name}")
            master_doc.add_heading(name, level=2)
            
            ext_files = process_folder(path, allowed_extensions)
            if ext_files:
                add_files_to_document(master_doc, ext_files, github_base_url)
            else:
                p = master_doc.add_paragraph("No files found or path does not exist.")
                p.italic = True
    
    # Save Master Document
    master_filename = "Full_Master_Overview.docx"
    master_save_path = os.path.join(output_folder, master_filename)
    master_doc.save(master_save_path)
    print(f"--- Master Overview saved to: {master_save_path} ---")

    print("--- All done! ---")

if __name__ == "__main__":
    generate_weekly_reports()